"""
Generates docs/Alphabot_User_Guide.docx from the markdown source.
Run: python3 docs/generate_user_guide.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

# ── Colour palette ────────────────────────────────────────────────────────────
C_INDIGO   = RGBColor(0x3B, 0x45, 0xDE)   # Platform Manager heading
C_SLATE    = RGBColor(0x47, 0x5C, 0x6F)   # Platform Admin heading
C_EMERALD  = RGBColor(0x05, 0x97, 0x69)   # Tenant Admin heading
C_AMBER    = RGBColor(0xD9, 0x77, 0x06)   # Agent heading
C_PURPLE   = RGBColor(0x7C, 0x3A, 0xED)   # Testing section
C_BLACK    = RGBColor(0x1A, 0x1A, 0x2E)
C_GREY     = RGBColor(0x6B, 0x72, 0x80)
C_LIGHTBG  = RGBColor(0xF3, 0xF4, 0xF6)
C_BORDER   = RGBColor(0xE5, 0xE7, 0xEB)
C_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)

SECTION_COLORS = {
    1: C_INDIGO,
    2: C_SLATE,
    3: C_EMERALD,
    4: C_AMBER,
    5: C_PURPLE,
}

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, sides=('top','bottom','left','right'), color='E5E7EB', sz='4'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def para_spacing(para, before=0, after=6):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)

def add_horizontal_rule(doc, color='E5E7EB'):
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pb   = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pb.append(bot)
    pPr.append(pb)
    para_spacing(para, 0, 0)
    return para

def add_run_bold(para, text, color=None, size=None):
    run = para.add_run(text)
    run.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run

def inline_bold(para, text):
    """Parse **bold** within text and add runs."""
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = para.add_run(part)
        if i % 2 == 1:
            run.bold = True

def inline_code(para, text):
    """Parse `code` within text and add styled runs."""
    parts = re.split(r'`(.*?)`', text)
    for i, part in enumerate(parts):
        run = para.add_run(part)
        if i % 2 == 1:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

def add_callout(doc, text, bg_hex='EEF2FF', border_hex='6366F1', icon='ℹ️'):
    """Indented callout box for notes/tips."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    set_cell_border(cell, color=border_hex, sz='12')
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    inline_bold(p, f'{icon}  {text}')
    doc.add_paragraph()

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, '1E293B')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(code_text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    doc.add_paragraph()

def add_screenshot_placeholder(doc, description):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, 'F9FAFB')
    set_cell_border(cell, color='D1D5DB', sz='6')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    r1 = p.add_run('📸  ')
    r1.font.size = Pt(18)
    r2 = p.add_run('SCREENSHOT\n')
    r2.bold = True
    r2.font.size  = Pt(11)
    r2.font.color.rgb = C_GREY
    r3 = p.add_run(description)
    r3.font.size  = Pt(9)
    r3.font.color.rgb = C_GREY
    doc.add_paragraph()

# ── Document setup ────────────────────────────────────────────────────────────

def setup_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default paragraph font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = C_BLACK

    return doc

# ── Cover page ────────────────────────────────────────────────────────────────

def add_cover(doc):
    # Big coloured banner
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, '3B45DE')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(32)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run('Alphabot')
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = C_WHITE
    r.font.name = 'Calibri'

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(32)
    r2 = p2.add_run('Complete User Guide')
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(0xC7, 0xD2, 0xFE)
    r2.font.name = 'Calibri'

    doc.add_paragraph()

    # Subtitle
    sub = doc.add_paragraph('WhatsApp AI Agent Platform')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(4)
    run = sub.runs[0]
    run.font.size = Pt(13)
    run.font.color.rgb = C_GREY

    ver = doc.add_paragraph('Version 1.0  ·  July 2026')
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = ver.runs[0]
    run2.font.size = Pt(9)
    run2.font.color.rgb = C_GREY

    doc.add_paragraph()
    add_horizontal_rule(doc)
    doc.add_paragraph()

    # Persona quick-reference
    personas = [
        ('Platform Manager', 'Full platform control — all clients, AI models, billing', '3B45DE'),
        ('Platform Admin',   'Read-only view of platform console',                      '475C6F'),
        ('Tenant Admin',     'Client workspace — bot config, team, orders, campaigns',  '059769'),
        ('Agent',            'Handle escalations and customer replies',                  'D97706'),
    ]

    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = 'Table Grid'
    header_cells = tbl2.rows[0].cells
    for cell, txt, width in zip(header_cells, ['Persona', 'Role', 'Section'], [4, 9, 2]):
        set_cell_bg(cell, '1E293B')
        p = cell.paragraphs[0]
        r = p.add_run(txt)
        r.bold = True
        r.font.color.rgb = C_WHITE
        r.font.size = Pt(9)

    for name, desc, color in personas:
        row = tbl2.add_row()
        set_cell_bg(row.cells[0], color)
        r0 = row.cells[0].paragraphs[0].add_run(name)
        r0.bold = True
        r0.font.color.rgb = C_WHITE
        r0.font.size = Pt(9)

        row.cells[1].paragraphs[0].add_run(desc).font.size = Pt(9)

        sec_num = personas.index((name, desc, color)) + 1
        row.cells[2].paragraphs[0].add_run(f'Section {sec_num}').font.size = Pt(9)

    doc.add_paragraph()
    doc.add_page_break()

# ── Section heading banner ────────────────────────────────────────────────────

def rgb_to_hex(color: RGBColor) -> str:
    return '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2])

def add_section_banner(doc, section_num, title, subtitle, color: RGBColor):
    hex_c = rgb_to_hex(color)
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, hex_c)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f'SECTION {section_num}  ')
    r1.font.size  = Pt(8)
    r1.font.color.rgb = C_WHITE
    r1.bold = True

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(title)
    r2.bold = True
    r2.font.size  = Pt(20)
    r2.font.color.rgb = C_WHITE
    r2.font.name  = 'Calibri'

    p3 = cell.add_paragraph()
    p3.paragraph_format.space_after = Pt(10)
    r3 = p3.add_run(subtitle)
    r3.font.size  = Pt(10)
    r3.font.color.rgb = RGBColor(0xE0, 0xE7, 0xFF)
    doc.add_paragraph()

# ── Heading helpers ───────────────────────────────────────────────────────────

def h2(doc, text, color=C_BLACK):
    p = doc.add_paragraph()
    para_spacing(p, before=14, after=4)
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(14)
    r.font.color.rgb = color
    add_horizontal_rule(doc)
    return p

def h3(doc, text, color=C_BLACK):
    p = doc.add_paragraph()
    para_spacing(p, before=10, after=3)
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(11.5)
    r.font.color.rgb = color
    return p

def h4(doc, text, color=C_GREY):
    p = doc.add_paragraph()
    para_spacing(p, before=8, after=2)
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(10)
    r.font.color.rgb = color
    return p

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    para_spacing(p, before=0, after=5)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    inline_bold(p, text)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    inline_bold(p, text)
    return p

def numbered(doc, text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    inline_bold(p, text)
    return p

# ── Generic table builder ─────────────────────────────────────────────────────

def add_table(doc, headers, rows, col_widths=None, header_bg='1E293B'):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = tbl.rows[0].cells
    for i, (cell, txt) in enumerate(zip(hdr, headers)):
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = C_WHITE
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.add_row()
        bg = 'FFFFFF' if ri % 2 == 0 else 'F9FAFB'
        for ci, (cell, txt) in enumerate(zip(row.cells, row_data)):
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            inline_bold(p, str(txt))
            for run in p.runs:
                run.font.size = Pt(9)

    if col_widths:
        for row in tbl.rows:
            for cell, w in zip(row.cells, col_widths):
                cell.width = Cm(w)

    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — Platform Manager
# ══════════════════════════════════════════════════════════════════════════════

def section_platform_manager(doc):
    add_section_banner(doc, 1, 'Platform Manager',
        'Full control over all clients, AI models, voice providers, billing and team.',
        C_INDIGO)

    # 1.1 Logging In
    h2(doc, '1.1  Logging In', C_INDIGO)
    numbered(doc, 'Go to the Alphabot web URL provided by your administrator.')
    numbered(doc, 'Enter your email and password on the login screen.')
    numbered(doc, 'You will be directed to the **Platform Console** — URL ends in `/platform/clients`.')
    add_screenshot_placeholder(doc, 'Login screen — email and password fields with Sign In button')
    add_callout(doc,
        'If you land on the client dashboard instead of the Platform Console, your account '
        'may not have a platform role assigned. Contact your system administrator.',
        bg_hex='FEF9C3', border_hex='CA8A04', icon='⚠️')

    # 1.2 Platform Console Overview
    h2(doc, '1.2  Platform Console Overview', C_INDIGO)
    body(doc, 'After login you arrive at the **Clients** page — the default landing for all platform users. '
              'The top navigation bar contains:')
    add_table(doc,
        ['Nav Item', 'What It Does'],
        [
            ('Clients',          'View and manage all onboarded client accounts'),
            ('Products',         'Manage the bot product catalog'),
            ('Guardrails',       'Set platform-wide content filters for all bots'),
            ('AI Models',        'Configure LLM providers and API keys'),
            ('Voice Providers',  'Set up telephony, STT, and TTS providers'),
            ('Analytics',        'View usage metrics across all clients'),
            ('Billing',          'Monitor subscriptions, trials, and revenue'),
            ('Notifications',    'Configure platform-level notification settings'),
            ('Settings',         'Platform-wide configuration'),
            ('Team',             'Manage platform users (Admin and Manager accounts)'),
        ],
        col_widths=[5, 11],
    )
    add_screenshot_placeholder(doc, 'Platform Console — top navigation bar with all menu items visible')

    # 1.3 Managing Clients
    h2(doc, '1.3  Managing Clients', C_INDIGO)

    h3(doc, '1.3.1  Viewing All Clients')
    body(doc, 'Navigate to **Clients** in the top nav. You will see a table of all onboarded client accounts.')
    add_table(doc,
        ['Column', 'Description'],
        [
            ('Client Name',      'Company or brand name'),
            ('Plan',             'Trial / Starter / Pro / Enterprise'),
            ('Status',           'Active or Suspended'),
            ('Active Bots',      'Number of bots currently running'),
            ('Joined',           'Onboarding date'),
        ],
        col_widths=[5, 11],
    )
    add_screenshot_placeholder(doc, 'Clients list page — table with plan badges and status indicators')
    body(doc, '**Filtering clients:** Use the search box (top-right) to filter by name, use the plan '
              'dropdown or status filter to narrow results.')
    add_screenshot_placeholder(doc, 'Clients page — search box active with filtered results visible')

    h3(doc, '1.3.2  Onboarding a New Client')
    numbered(doc, 'Click **+ New Client** (top-right of Clients page).')
    numbered(doc, 'Fill in the form:')
    add_table(doc,
        ['Field', 'Description'],
        [
            ('Company Name', 'The client\'s business name'),
            ('Admin Email',  'Primary admin user\'s email — they receive an invite'),
            ('Plan',         'Select Trial, Starter, Pro, or Enterprise'),
            ('Products',     'Tick the bot products this client gets access to'),
        ],
        col_widths=[5, 11],
    )
    numbered(doc, 'Click **Create Client**.')
    numbered(doc, 'The client admin will receive an email invitation to set up their account.')
    add_screenshot_placeholder(doc, 'New Client form — all fields filled in')
    add_screenshot_placeholder(doc, 'Success confirmation toast after client is created')

    h3(doc, '1.3.3  Viewing a Client\'s Details')
    body(doc, 'Click any client row to open the **Client Detail** page. You will see company '
              'information, active WhatsApp numbers, team members, and recent activity.')
    add_screenshot_placeholder(doc, 'Client detail page — company info, WhatsApp numbers, and team members listed')

    # 1.4 Product Catalog
    h2(doc, '1.4  Managing the Product Catalog', C_INDIGO)
    body(doc, 'Navigate to **Products** in the top nav. Products are bot types '
              '(e.g., "Water Purifier Sales Bot"). Each product has a unique slug, display name, and description.')
    h3(doc, 'Adding a New Product')
    numbered(doc, 'Click **+ New Product**.')
    numbered(doc, 'Enter the **Slug** (lowercase, no spaces — e.g., `water-purifier`), **Name**, and **Description**.')
    numbered(doc, 'Click **Save**.')
    add_screenshot_placeholder(doc, 'Products list page with existing products shown')
    add_screenshot_placeholder(doc, 'New Product form — slug, name, description fields')

    # 1.5 Platform Guardrails
    h2(doc, '1.5  Platform Guardrails', C_INDIGO)
    body(doc, 'Navigate to **Guardrails**. These rules apply to **every bot on the platform** '
              'regardless of what individual clients configure.')
    add_table(doc,
        ['Setting', 'Description'],
        [
            ('Blocked Topics',       'Keywords or topic categories the bot must never discuss'),
            ('Max Message Length',   'Maximum characters per bot reply'),
            ('Language Restrictions','Allowed response languages'),
            ('Escalation Keywords',  'Words that immediately trigger human handoff'),
            ('Profanity Filter',     'Enable/disable automatic profanity blocking'),
        ],
        col_widths=[5, 11],
    )
    numbered(doc, 'Edit any field directly on the page.')
    numbered(doc, 'Click **Save Changes**. Changes take effect within 60 seconds.')
    add_screenshot_placeholder(doc, 'Platform Guardrails page — blocked topics and escalation keywords fields')
    add_callout(doc, 'These rules cannot be overridden by any client. They are the platform-wide minimum safety standard.',
                bg_hex='FEF2F2', border_hex='EF4444', icon='🛡️')

    # 1.6 AI Models
    h2(doc, '1.6  AI Model Configuration', C_INDIGO)
    body(doc, 'Navigate to **AI Models**. Configure which LLM providers are available and set API keys.')
    h3(doc, 'Adding a New LLM Configuration')
    numbered(doc, 'Click **+ Add Config**.')
    numbered(doc, 'Select a **Provider** (Anthropic, OpenAI, etc.).')
    numbered(doc, 'Enter the **API Key**.')
    numbered(doc, 'Select the **Model** (e.g., `claude-sonnet-4-6`).')
    numbered(doc, 'Optionally enter a **Base URL** (for self-hosted or proxy endpoints).')
    numbered(doc, 'Assign to a **Tenant** — leave blank for global default.')
    numbered(doc, 'Click **Save**.')
    add_screenshot_placeholder(doc, 'AI Models page — list of model configs with provider, model name, tenant columns')
    add_screenshot_placeholder(doc, 'Add LLM Config form — provider dropdown, API key, model, tenant fields')

    # 1.7 Voice Providers
    h2(doc, '1.7  Voice Providers', C_INDIGO)
    body(doc, 'Navigate to **Voice Providers**. This manages the three layers of the voice call stack.')
    add_table(doc,
        ['Component', 'Purpose', 'Example Providers'],
        [
            ('Telephony', 'Handles actual call routing',             'Exotel, Twilio'),
            ('STT',       'Transcribes caller\'s speech',            'Sarvam AI, Deepgram, Azure'),
            ('TTS',       'Converts bot reply to audio',             'ExotelSay, Google TTS, Sarvam TTS'),
        ],
        col_widths=[4, 7, 5],
    )
    body(doc, 'The **Default stack cost** card shows estimated cost per minute in ₹.')
    add_screenshot_placeholder(doc, 'Voice Providers page — three sections (Telephony, STT, TTS) with provider cards')
    h3(doc, 'Enabling a Provider')
    numbered(doc, 'Click the **Edit (pencil)** icon on any provider card.')
    numbered(doc, 'Enter credentials (API key, Account SID, etc.).')
    numbered(doc, 'Toggle **Enabled** to ON.')
    numbered(doc, 'Optionally toggle **Set as Default**.')
    numbered(doc, 'Click **Save**.')
    add_screenshot_placeholder(doc, 'Edit credentials modal for a telephony provider — credential fields visible')
    h3(doc, 'Recommended Setup Order')
    numbered(doc, 'Telephony (Exotel or Twilio) → Enter credentials → Enable → Set as Default')
    numbered(doc, 'STT (Sarvam or Deepgram) → Enter credentials → Enable → Set as Default')
    numbered(doc, 'TTS — ExotelSay is free and included → Enable → Set as Default')
    numbered(doc, 'In each bot\'s **Guardrails** settings → enable Voice → set greeting and language')

    # 1.8 Analytics
    h2(doc, '1.8  Platform Analytics', C_INDIGO)
    body(doc, 'Navigate to **Analytics** to view platform-wide metrics across all clients.')
    add_table(doc,
        ['Metric', 'Description'],
        [
            ('Total Messages',    'All messages processed across all bots'),
            ('Resolution Rate',   '% of conversations resolved without escalation'),
            ('Escalation Rate',   '% of conversations escalated to human agents'),
            ('Avg Response Time', 'Average bot reply latency'),
            ('Active Tenants',    'Clients with activity in the period'),
            ('Top Products',      'Bot products with the highest volume'),
        ],
        col_widths=[5, 11],
    )
    add_screenshot_placeholder(doc, 'Platform Analytics page — metric cards at top, line charts below')
    add_screenshot_placeholder(doc, 'Date range picker open with custom range selected')

    # 1.9 Billing
    h2(doc, '1.9  Billing & Subscriptions', C_INDIGO)
    body(doc, 'Navigate to **Billing** to monitor all client subscriptions, trial expirations, '
              'monthly active usage, and AI token consumption.')
    add_screenshot_placeholder(doc, 'Billing page — client rows with plan, renewal date, trial status columns')
    add_callout(doc, 'Clients with expired trials or failed payments are highlighted in red. '
                'Click their row to view details and take action.', bg_hex='FEF2F2', border_hex='EF4444', icon='⚠️')

    # 1.10 Notifications
    h2(doc, '1.10  Notification Settings', C_INDIGO)
    body(doc, 'Navigate to **Notifications** to configure where platform-level alerts are sent.')
    add_table(doc,
        ['Alert Type', 'Description'],
        [
            ('New client signup',    'When a new client onboards'),
            ('Trial expiring',       'When a client\'s trial has < 7 days left'),
            ('Payment failure',      'When a subscription payment fails'),
            ('High escalation rate', 'When a client\'s escalation rate exceeds threshold'),
        ],
        col_widths=[6, 10],
    )
    body(doc, 'Enter email addresses or webhook URLs for each alert type, then click **Save**.')
    add_screenshot_placeholder(doc, 'Notifications settings — alert types with destination input fields')

    # 1.11 Team
    h2(doc, '1.11  Team Management (Platform Users)', C_INDIGO)
    body(doc, 'Navigate to **Team** to manage platform-level staff accounts.')
    h3(doc, 'Inviting a New Platform User')
    numbered(doc, 'Click **+ Invite User**.')
    numbered(doc, 'Enter their **email address**.')
    numbered(doc, 'Select their **role**: **Manager** (full access) or **Admin** (read-only).')
    numbered(doc, 'Click **Send Invite**. The user receives an email to set up their account.')
    add_screenshot_placeholder(doc, 'Team page — list of platform users with Manager/Admin role badges')
    add_screenshot_placeholder(doc, 'Invite User modal — email field and role dropdown')
    h3(doc, 'Removing a Team Member')
    numbered(doc, 'Find the user in the Team list.')
    numbered(doc, 'Click the **⋮ three-dot menu** on their row.')
    numbered(doc, 'Select **Remove User** and confirm.')
    add_screenshot_placeholder(doc, 'Three-dot menu open on a team member row — Remove User option visible')

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — Platform Admin
# ══════════════════════════════════════════════════════════════════════════════

def section_platform_admin(doc):
    add_section_banner(doc, 2, 'Platform Admin',
        'Read-only access to the Platform Console — view all data without making changes.',
        C_SLATE)

    h2(doc, '2.1  Logging In', C_SLATE)
    body(doc, 'Same process as Platform Manager. After login you land on the **Platform Console**.')
    add_screenshot_placeholder(doc, 'Platform Console landing page — Admin perspective')

    h2(doc, '2.2  What You Can View', C_SLATE)
    add_table(doc,
        ['Page', 'Your Access'],
        [
            ('Clients',          'View client list and details — no create/edit buttons'),
            ('Products',         'View product catalog'),
            ('Guardrails',       'View platform guardrail settings'),
            ('AI Models',        'View model configurations (API keys are masked)'),
            ('Voice Providers',  'View provider configs (credentials masked)'),
            ('Analytics',        'Full read access'),
            ('Billing',          'View subscription status'),
            ('Notifications',    'View notification destinations'),
            ('Team',             'View team members — cannot invite or remove'),
        ],
        col_widths=[5, 11],
    )
    add_screenshot_placeholder(doc, 'Admin view of Clients page — "+ New Client" button absent')
    add_callout(doc, 'Action buttons (+ New, Edit, Save, Delete) are not visible on your account. '
                'This is by design for the Admin role.', bg_hex='EFF6FF', border_hex='3B82F6', icon='ℹ️')

    h2(doc, '2.3  Viewing Client Details', C_SLATE)
    numbered(doc, 'Click any client in the Clients list.')
    numbered(doc, 'View subscription info, active bots, and team members (read-only).')
    add_screenshot_placeholder(doc, 'Client detail page in read-only Admin view')

    h2(doc, '2.4  Viewing Analytics', C_SLATE)
    body(doc, 'Navigate to **Analytics**. All charts and metrics are available. '
              'Use the date range filter to explore specific periods.')
    add_screenshot_placeholder(doc, 'Analytics page — all charts visible, no edit controls')

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — Tenant Admin / Supervisor
# ══════════════════════════════════════════════════════════════════════════════

def section_tenant_admin(doc):
    add_section_banner(doc, 3, 'Tenant Admin / Supervisor',
        'Manage your team, configure your bot, handle escalations, create orders, and run campaigns.',
        C_EMERALD)

    # 3.1 Login
    h2(doc, '3.1  Logging In', C_EMERALD)
    numbered(doc, 'Open the Alphabot invitation email from your account manager.')
    numbered(doc, 'Click the **Accept Invite** link.')
    numbered(doc, 'Set your password.')
    numbered(doc, 'You will be directed to your **Dashboard**.')
    add_screenshot_placeholder(doc, 'Accept Invite email with button highlighted')
    add_screenshot_placeholder(doc, 'Set Password screen')
    add_screenshot_placeholder(doc, 'Dashboard home page after first login')

    # 3.2 Dashboard
    h2(doc, '3.2  Dashboard Overview', C_EMERALD)
    body(doc, 'The Dashboard home shows key real-time metrics for your workspace:')
    add_table(doc,
        ['Card', 'Description'],
        [
            ('Open Conversations', 'Conversations currently being handled by the bot'),
            ('Pending Escalations','Conversations waiting for a human agent'),
            ('Resolved Today',     'Conversations marked resolved in the last 24 hours'),
            ('Response Rate',      '% of customer messages the bot replied to'),
        ],
        col_widths=[5, 11],
    )
    add_screenshot_placeholder(doc, 'Dashboard home — four metric cards at the top')
    body(doc, 'The left sidebar navigation:')
    add_table(doc,
        ['Nav Item', 'What It Does'],
        [
            ('Dashboard',      'Home metrics overview'),
            ('Conversations',  'All chat conversations'),
            ('Escalations',    'Conversations needing human attention'),
            ('Team',           'Manage agents and supervisors'),
            ('Knowledge Base', 'Q&A content that powers your bot'),
            ('Orders',         'Create and manage customer orders'),
            ('Follow-ups',     'Automated follow-up message sequences'),
            ('Campaigns',      'Bulk WhatsApp or voice outreach'),
            ('Voice',          'Voice call management'),
            ('Analytics',      'Your workspace analytics'),
            ('Billing',        'View subscription and usage'),
            ('Settings',       'Configure bots, WhatsApp numbers, workspace'),
            ('Guardrails',     'Configure content rules for your bots'),
            ('AI Models',      'Configure LLM for your workspace'),
        ],
        col_widths=[5, 11],
    )
    add_screenshot_placeholder(doc, 'Left sidebar — all navigation items visible')

    # 3.3 Conversations
    h2(doc, '3.3  Conversations', C_EMERALD)
    h3(doc, '3.3.1  Viewing All Conversations')
    body(doc, 'Navigate to **Conversations**. You see a list of all customer conversations with '
              'customer name/phone, last message preview, bot name, status, and timestamp.')
    add_screenshot_placeholder(doc, 'Conversations list — full page with filters at top and conversation rows')
    body(doc, '**Filtering:**')
    bullet(doc, '**By Bot** — select a product from the dropdown')
    bullet(doc, '**By Status** — toggle Open / Escalated / Resolved tabs')
    bullet(doc, '**Search** — type a phone number or customer name')
    add_screenshot_placeholder(doc, 'Conversations filtered by Escalated status')

    h3(doc, '3.3.2  Viewing a Conversation Thread')
    body(doc, 'Click any conversation to open the full message thread.')
    add_screenshot_placeholder(doc, 'Conversation detail — message thread with customer info panel on right')
    body(doc, 'Actions available:')
    add_table(doc,
        ['Action', 'Description'],
        [
            ('Escalate to Agent', 'Flags this conversation for human follow-up'),
            ('Mark Resolved',     'Closes the conversation'),
            ('Assign to Agent',   'Assign to a specific team member'),
            ('Send Message',      'Type and send a WhatsApp message directly'),
        ],
        col_widths=[5, 11],
    )
    add_screenshot_placeholder(doc, 'Conversation detail — action buttons (Escalate, Resolve, Assign, Send) visible')

    # 3.4 Escalations
    h2(doc, '3.4  Escalations', C_EMERALD)
    body(doc, 'Navigate to **Escalations**. This shows all conversations flagged for human attention '
              '— either by the bot (low confidence) or manually. Each card shows customer info, '
              'last messages, time since escalation, and assignment status.')
    add_screenshot_placeholder(doc, 'Escalations page — list of escalation cards with timestamps')
    h3(doc, 'Claiming an Escalation')
    numbered(doc, 'Click **Claim** on an escalation card.')
    numbered(doc, 'The conversation is assigned to you.')
    numbered(doc, 'Click the conversation to open the thread and respond.')
    numbered(doc, 'When resolved, click **Mark Resolved**.')
    add_screenshot_placeholder(doc, 'Escalation card with Claim button')
    add_screenshot_placeholder(doc, 'Escalation conversation open — agent typing a reply')

    # 3.5 Team
    h2(doc, '3.5  Team Management', C_EMERALD)
    h3(doc, '3.5.1  Inviting a New Team Member')
    numbered(doc, 'Click **+ Invite Member**.')
    numbered(doc, 'Enter their **email address**.')
    numbered(doc, 'Select their **role**:')
    add_table(doc,
        ['Role', 'Permissions'],
        [
            ('Admin',      'Full access — configure settings, manage team, create orders'),
            ('Supervisor', 'View and manage conversations, escalations, analytics'),
            ('Agent',      'View conversations and handle escalations only'),
        ],
        col_widths=[4, 12],
    )
    numbered(doc, 'Click **Send Invite**. Team member receives an email invitation.')
    add_screenshot_placeholder(doc, 'Team page — members with Admin, Supervisor, Agent role badges')
    add_screenshot_placeholder(doc, 'Invite Member modal — email field and role dropdown')

    h3(doc, '3.5.2  Removing a Team Member')
    numbered(doc, 'Find the member in the Team list.')
    numbered(doc, 'Click the **⋮ menu** on their row.')
    numbered(doc, 'Select **Remove** and confirm.')
    add_screenshot_placeholder(doc, 'Three-dot menu on a team member row — Remove option')

    # 3.6 Knowledge Base
    h2(doc, '3.6  Knowledge Base', C_EMERALD)
    body(doc, 'The Knowledge Base is the set of Q&A pairs your bot uses to answer customer '
              'questions accurately. KB entries are grouped into **Collections** '
              '(e.g., "Product FAQs", "Pricing", "Warranty").')
    h3(doc, '3.6.1  Adding a New KB Entry')
    numbered(doc, 'Click a collection to open it.')
    numbered(doc, 'Click **+ Add Entry**.')
    numbered(doc, 'Fill in the **Question** (customer query) and **Answer** (what the bot should say).')
    numbered(doc, 'Click **Save**. Changes are live within 60 seconds.')
    add_screenshot_placeholder(doc, 'Knowledge Base — collections list with entry counts')
    add_screenshot_placeholder(doc, 'KB collection detail — list of Q&A entries with edit and delete icons')
    add_screenshot_placeholder(doc, 'Add Entry form — Question and Answer fields')

    h3(doc, '3.6.2  Creating a New Collection')
    numbered(doc, 'On the KB main page, click **+ New Collection**.')
    numbered(doc, 'Enter a **Name** and optional **Description**.')
    numbered(doc, 'Click **Create**.')
    add_screenshot_placeholder(doc, 'New Collection form — name and description fields')

    # 3.7 Orders
    h2(doc, '3.7  Orders', C_EMERALD)
    h3(doc, '3.7.1  Creating a New Order')
    numbered(doc, 'Click **+ New Order**.')
    numbered(doc, 'Fill in the order form:')
    add_table(doc,
        ['Field', 'Description'],
        [
            ('Customer Name',   'The buyer\'s name'),
            ('Phone Number',    'WhatsApp number with country code (e.g., 919876543210)'),
            ('Product',         'Select from your product catalog'),
            ('Quantity',        'Number of units'),
            ('Amount',          'Total order value in ₹'),
            ('Payment Method',  'Razorpay / PhonePe / Manual'),
            ('Notes',           'Internal notes — not sent to customer'),
        ],
        col_widths=[5, 11],
    )
    numbered(doc, 'Click **Create Order**. If a payment method is selected, a payment link is '
                  'automatically generated and sent to the customer via WhatsApp.')
    add_screenshot_placeholder(doc, 'Orders list — rows with status badges (Pending, Paid, Delivered)')
    add_screenshot_placeholder(doc, 'New Order form — all fields visible')
    add_screenshot_placeholder(doc, 'Payment link confirmation toast after order creation')

    h3(doc, '3.7.2  Order Status Reference')
    add_table(doc,
        ['Status', 'Meaning'],
        [
            ('Pending',    'Order created, awaiting payment'),
            ('Paid',       'Payment received'),
            ('Processing', 'Order in fulfilment'),
            ('Delivered',  'Order fulfilled'),
            ('Cancelled',  'Order cancelled'),
        ],
        col_widths=[4, 12],
    )
    add_screenshot_placeholder(doc, 'Order detail page — payment status, customer info, linked conversation')

    # 3.8 Follow-ups
    h2(doc, '3.8  Follow-ups', C_EMERALD)
    body(doc, 'Follow-ups are automated WhatsApp messages sent to customers at scheduled intervals '
              'after a conversation or order event.')
    numbered(doc, 'Click **+ New Follow-up**.')
    numbered(doc, 'Set the **Trigger** (e.g., Order Created, Conversation Ended).')
    numbered(doc, 'Set the **Delay** (e.g., 2 hours, 1 day).')
    numbered(doc, 'Write the **Message** to send.')
    numbered(doc, 'Click **Save**.')
    add_screenshot_placeholder(doc, 'Follow-ups list — active sequences with trigger and delay')
    add_screenshot_placeholder(doc, 'New Follow-up form — trigger, delay, message fields')

    # 3.9 Campaigns
    h2(doc, '3.9  Campaigns', C_EMERALD)
    body(doc, 'Campaigns let you send **bulk outreach** to a list of customers via WhatsApp, '
              'voice calls, or both.')
    h3(doc, '3.9.1  Creating a Campaign')
    numbered(doc, 'Click **+ New Campaign**.')
    numbered(doc, 'Fill in the campaign form:')
    add_table(doc,
        ['Field', 'Description'],
        [
            ('Campaign Name',     'Internal name for tracking'),
            ('Channel',           'WhatsApp / Voice / Both'),
            ('Message Template',  'The WhatsApp message body'),
            ('Schedule',          'Send now or schedule for a future date/time'),
            ('Contacts',          'Upload a CSV or paste phone numbers'),
        ],
        col_widths=[5, 11],
    )
    numbered(doc, 'Click **Launch Campaign**.')
    add_screenshot_placeholder(doc, 'Campaigns list — channel badges (WhatsApp, Voice, Both) and status')
    add_screenshot_placeholder(doc, 'New Campaign form — channel selector and message template')
    add_screenshot_placeholder(doc, 'Contacts section — CSV upload button and preview of contacts')

    h3(doc, '3.9.2  Campaign Stats')
    add_table(doc,
        ['Stat', 'Description'],
        [
            ('Sent',            'Messages successfully delivered'),
            ('Replied',         'Customers who replied'),
            ('Failed',          'Messages that could not be delivered'),
            ('Calls Made',      '(Voice) Calls dialled'),
            ('Calls Answered',  '(Voice) Calls where customer picked up'),
            ('Voicemails Left', '(Voice) Calls that went to voicemail'),
        ],
        col_widths=[5, 11],
    )
    add_screenshot_placeholder(doc, 'Campaign detail — stats cards and per-contact status list')

    # 3.10 Settings
    h2(doc, '3.10  Bot Configuration (Settings)', C_EMERALD)
    h3(doc, '3.10.1  Workspace Settings')
    body(doc, 'Navigate to **Settings → Workspace** tab. Set your workspace name and timezone.')
    add_screenshot_placeholder(doc, 'Settings — Workspace tab with name and timezone fields')

    h3(doc, '3.10.2  Bot Configuration')
    body(doc, 'Click the **Bot Config** tab. For each bot product:')
    add_table(doc,
        ['Setting', 'Description'],
        [
            ('System Prompt',        'Instructions that define your bot\'s personality and scope'),
            ('AI Model',             'Which LLM model to use'),
            ('Confidence Threshold', 'Score below which the bot escalates (0.0–1.0, default 0.7)'),
            ('Escalation Triggers',  'Keywords that immediately escalate to a human'),
            ('Escalation Policy',    'Who to notify and how when an escalation fires'),
        ],
        col_widths=[5, 11],
    )
    add_screenshot_placeholder(doc, 'Bot Config form — system prompt, confidence threshold slider, escalation triggers')
    add_callout(doc, 'The system prompt is the most powerful tool for shaping bot behaviour. '
                'Be specific about what topics the bot should and should not discuss.',
                bg_hex='F0FDF4', border_hex='22C55E', icon='💡')

    h3(doc, '3.10.3  WhatsApp Number Setup')
    body(doc, 'Click the **WhatsApp Setup** tab to link your WhatsApp Business number.')
    h4(doc, 'For Twilio:')
    numbered(doc, 'Enter your Twilio **Account SID** and **Auth Token**.')
    numbered(doc, 'Enter your **Twilio WhatsApp Number** (e.g., `+14155238886`).')
    numbered(doc, 'Click **Save**.')
    numbered(doc, 'Copy the **Webhook URL** shown and paste it into Twilio Console → Sandbox settings.')
    add_screenshot_placeholder(doc, 'WhatsApp Setup — Twilio credentials form and webhook URL field')
    h4(doc, 'For Meta Cloud API:')
    numbered(doc, 'Enter your **Phone Number ID** and **Access Token** from Meta Business Manager.')
    numbered(doc, 'Enter the **Verify Token** you set in Meta.')
    numbered(doc, 'Click **Save**, then copy the **Webhook URL** into Meta\'s App Dashboard.')
    add_screenshot_placeholder(doc, 'WhatsApp Setup — Meta Cloud API form')

    # 3.11 Guardrails
    h2(doc, '3.11  Guardrails (Bot-level)', C_EMERALD)
    body(doc, 'Navigate to **Guardrails**. These rules apply only to your bots, in addition to platform rules.')
    add_table(doc,
        ['Setting', 'Description'],
        [
            ('Blocked Topics',    'Topics your bot should refuse to discuss'),
            ('Allowed Languages', 'Languages the bot should reply in'),
            ('Max Reply Length',  'Maximum characters per reply'),
            ('Voice Enabled',     'Whether this bot can receive/make calls'),
            ('Voice Greeting',    'Opening message for voice calls'),
            ('Voice Language',    'Language for text-to-speech output'),
        ],
        col_widths=[5, 11],
    )
    body(doc, 'After changes, click **Save Guardrails**. Changes are live within 60 seconds.')
    add_screenshot_placeholder(doc, 'Guardrails page — blocked topics, language settings, voice toggle')

    # 3.12 Analytics
    h2(doc, '3.12  Analytics', C_EMERALD)
    body(doc, 'Navigate to **Analytics** to view your workspace-specific metrics:')
    add_table(doc,
        ['Metric', 'Description'],
        [
            ('Message Volume',         'Total messages over time (chart)'),
            ('Resolution Rate',        '% of conversations resolved by bot'),
            ('Escalation Rate',        '% escalated to humans'),
            ('Avg Response Time',      'How fast the bot replies'),
            ('Top Questions',          'Most common customer queries'),
            ('Confidence Distribution','Distribution of bot confidence scores'),
        ],
        col_widths=[5, 11],
    )
    add_screenshot_placeholder(doc, 'Analytics page — line chart for message volume and metric cards')
    add_screenshot_placeholder(doc, 'Top Questions table — question text and frequency count')

    # 3.13 Billing
    h2(doc, '3.13  Billing', C_EMERALD)
    body(doc, 'Navigate to **Billing** to view your current plan, renewal date, '
              'monthly message usage, AI token usage, and payment history.')
    add_screenshot_placeholder(doc, 'Billing page — plan card, usage progress bar, token usage card')

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — Agent
# ══════════════════════════════════════════════════════════════════════════════

def section_agent(doc):
    add_section_banner(doc, 4, 'Agent',
        'Handle escalated conversations and reply to customers on behalf of your team.',
        C_AMBER)

    h2(doc, '4.1  Logging In', C_AMBER)
    numbered(doc, 'Open the invitation email from your manager/admin.')
    numbered(doc, 'Click **Accept Invite**.')
    numbered(doc, 'Set your password.')
    numbered(doc, 'You will be directed to the **Dashboard**.')
    add_screenshot_placeholder(doc, 'Agent dashboard — simplified view with fewer sidebar items than Admin')

    h2(doc, '4.2  What You Can Access', C_AMBER)
    body(doc, 'As an Agent, your sidebar shows a limited set of pages:')
    add_table(doc,
        ['Page', 'Your Access'],
        [
            ('Dashboard',    'View metrics (read-only)'),
            ('Conversations','View all conversations'),
            ('Escalations',  '⭐ View and claim escalations — your primary workspace'),
            ('Analytics',    'View metrics (read-only)'),
        ],
        col_widths=[5, 11],
    )
    add_screenshot_placeholder(doc, 'Agent sidebar — only Dashboard, Conversations, Escalations, Analytics visible')

    h2(doc, '4.3  Handling Escalations — Step-by-Step', C_AMBER)
    body(doc, 'This is your **primary workflow**. Follow these five steps for every escalation.')

    h3(doc, 'Step 1 — Navigate to Escalations')
    body(doc, 'Click **Escalations** in the sidebar. You see all conversations that need human '
              'attention, sorted oldest first (most urgent at top).')
    add_screenshot_placeholder(doc, 'Escalations queue — cards sorted by time waiting, oldest at top with amber timestamps')

    h3(doc, 'Step 2 — Claim a Conversation')
    numbered(doc, 'Review the escalation card: customer name/number, why the bot escalated, '
                  'and the last few messages.')
    numbered(doc, 'Click **Claim** to assign it to yourself. It is removed from other agents\' queues.')
    add_screenshot_placeholder(doc, 'Escalation card — customer info, escalation reason, last messages, Claim button')

    h3(doc, 'Step 3 — Open the Conversation')
    numbered(doc, 'After claiming, click **Open Conversation** to see the full message thread.')
    numbered(doc, 'Read the history to understand what the customer needs.')
    add_screenshot_placeholder(doc, 'Conversation thread — full history, customer on left, bot on right, reply area at bottom')

    h3(doc, 'Step 4 — Reply to the Customer')
    numbered(doc, 'Type your response in the **message box** at the bottom.')
    numbered(doc, 'Click **Send** (or press Enter).')
    numbered(doc, 'Your message is sent via WhatsApp to the customer immediately.')
    add_screenshot_placeholder(doc, 'Message box at bottom — agent typing a reply')
    add_screenshot_placeholder(doc, 'Agent reply appearing in the thread as a new message bubble')

    h3(doc, 'Step 5 — Resolve the Conversation')
    numbered(doc, 'When the issue is resolved, click **Mark Resolved** (top-right).')
    numbered(doc, 'The conversation moves to the Resolved tab.')
    numbered(doc, 'The bot resumes normal handling for future messages from this customer.')
    add_screenshot_placeholder(doc, 'Mark Resolved button at top right of conversation view')
    add_screenshot_placeholder(doc, 'Conversation moved to Resolved — green Resolved badge')

    h2(doc, '4.4  Viewing Conversations', C_AMBER)
    body(doc, 'Navigate to **Conversations** to browse all conversations (not just escalations). '
              'Filter by status: Open, Escalated, or Resolved.')
    add_screenshot_placeholder(doc, 'Conversations list from agent view — status filter tabs visible')

    h2(doc, '4.5  Viewing Analytics', C_AMBER)
    body(doc, 'Navigate to **Analytics** to see workspace-level metrics. This view is read-only — '
              'you cannot change any settings from here.')
    add_screenshot_placeholder(doc, 'Analytics page — read-only view from agent perspective')

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — Twilio Sandbox Testing
# ══════════════════════════════════════════════════════════════════════════════

def section_testing(doc):
    add_section_banner(doc, 5, 'Testing with Twilio Sandbox',
        'Add new test devices and verify your bot before going live with a production number.',
        C_PURPLE)

    h2(doc, '5.1  What is the Twilio Sandbox?', C_PURPLE)
    body(doc, 'The Twilio Sandbox is a shared test WhatsApp number provided by Twilio. It lets you '
              'test your bot without needing a verified WhatsApp Business number. '
              'Any device can send messages to the sandbox after a one-time opt-in step.')
    add_table(doc,
        ['Detail', 'Value'],
        [
            ('Sandbox Number', '+1 415 523 8886'),
            ('Join Code',      'join nearly-home'),
            ('Opt-in validity','Active as long as you message within ~72 hours'),
        ],
        col_widths=[5, 11],
    )

    h2(doc, '5.2  Adding a New Test Device', C_PURPLE)
    body(doc, 'Repeat this for **every phone number** that wants to test the bot.')
    h3(doc, 'Step 1')
    body(doc, 'Open **WhatsApp** on the test device.')
    h3(doc, 'Step 2')
    body(doc, 'Start a new chat with **+1 415 523 8886**.')
    add_screenshot_placeholder(doc, 'WhatsApp — new chat search showing +14155238886')
    h3(doc, 'Step 3')
    body(doc, 'Send this exact message:')
    add_code_block(doc, 'join nearly-home')
    add_screenshot_placeholder(doc, 'WhatsApp chat with sandbox number — "join nearly-home" sent')
    h3(doc, 'Step 4')
    body(doc, 'Twilio will reply with a confirmation:')
    add_callout(doc,
        '"You are now connected to the sandbox and can send and receive messages."',
        bg_hex='F0FDF4', border_hex='22C55E', icon='✅')
    add_screenshot_placeholder(doc, 'Twilio confirmation reply in WhatsApp')
    h3(doc, 'Step 5')
    body(doc, 'Send any message to start testing the bot:')
    add_code_block(doc, 'Hello')
    body(doc, 'The bot should reply within a few seconds.')
    add_screenshot_placeholder(doc, 'Test conversation — bot reply visible after "Hello" message')

    h2(doc, '5.3  Verifying the Webhook URL', C_PURPLE)
    body(doc, 'If the bot is not replying, verify the webhook is configured in Twilio.')
    numbered(doc, 'Log in to the **Twilio Console**.')
    numbered(doc, 'Go to **Messaging → Try it out → Send a WhatsApp message**.')
    numbered(doc, 'Click **Sandbox settings**.')
    numbered(doc, 'Confirm the **"WHEN A MESSAGE COMES IN"** field contains your Render API URL:')
    add_code_block(doc, 'https://your-app.onrender.com/api/webhook/twilio')
    numbered(doc, 'Click **Save**.')
    add_screenshot_placeholder(doc, 'Twilio Sandbox settings page — "When a message comes in" webhook URL field highlighted')

    h2(doc, '5.4  Troubleshooting', C_PURPLE)
    add_table(doc,
        ['Symptom', 'Likely Cause', 'Fix'],
        [
            ('Bot not replying at all',
             'Webhook URL wrong or Render app is down',
             'Check Twilio sandbox settings and Render dashboard'),
            ('Bot replies "could not find your account"',
             'whatsapp_numbers row missing or inactive',
             'Check Supabase → whatsapp_numbers table, ensure active = true'),
            ('Confidence score visible in message (e.g., CONFIDENCE:0.95)',
             'Running an older build',
             'Trigger a new deploy on Render — the bug is already fixed'),
            ('Test number not joining sandbox',
             '"join nearly-home" not sent or typo',
             'Re-send the exact join code from the test device'),
            ('Sandbox works but production number does not',
             'Production number not registered as WhatsApp Business sender',
             'Complete Twilio or Meta sender registration process'),
        ],
        col_widths=[5, 6, 5],
    )


# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX — Role Permissions Summary
# ══════════════════════════════════════════════════════════════════════════════

def section_appendix(doc):
    doc.add_page_break()
    p = doc.add_paragraph()
    para_spacing(p, before=0, after=8)
    r = p.add_run('Appendix: Role Permissions Summary')
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = C_BLACK

    add_table(doc,
        ['Feature', 'Plt Manager', 'Plt Admin', 'Tenant Admin', 'Supervisor', 'Agent'],
        [
            ('Platform Console',        '✅ Full',  '✅ Read',  '❌', '❌', '❌'),
            ('Manage Clients',          '✅',       '👁 View',  '❌', '❌', '❌'),
            ('Invite Platform Users',   '✅',       '❌',       '❌', '❌', '❌'),
            ('Platform Analytics',      '✅',       '✅',       '❌', '❌', '❌'),
            ('Platform Guardrails',     '✅',       '❌',       '❌', '❌', '❌'),
            ('Configure AI Models (Plt)','✅',      '❌',       '❌', '❌', '❌'),
            ('Configure Voice Providers','✅',      '❌',       '❌', '❌', '❌'),
            ('Client Dashboard',        '❌',       '❌',       '✅', '✅', '✅'),
            ('View Conversations',       '❌',      '❌',       '✅', '✅', '✅'),
            ('Handle Escalations',       '❌',      '❌',       '✅', '✅', '✅'),
            ('Send Messages to Customers','❌',     '❌',       '✅', '✅', '✅'),
            ('Invite Team Members',      '❌',      '❌',       '✅', '❌', '❌'),
            ('Remove Team Members',      '❌',      '❌',       '✅', '❌', '❌'),
            ('Manage Knowledge Base',    '❌',      '❌',       '✅', '👁 View','👁 View'),
            ('Create Orders',            '❌',      '❌',       '✅', '❌', '❌'),
            ('Create Campaigns',         '❌',      '❌',       '✅', '❌', '❌'),
            ('Configure Bot Settings',   '❌',      '❌',       '✅', '❌', '❌'),
            ('Bot-level Guardrails',     '❌',      '❌',       '✅', '❌', '❌'),
            ('View Own Analytics',       '❌',      '❌',       '✅', '✅', '✅'),
            ('View Billing',             '❌',      '❌',       '✅', '❌', '❌'),
        ],
        col_widths=[5.5, 2.5, 2.5, 2.5, 2.5, 2],
        header_bg='1E293B',
    )

    doc.add_paragraph()
    foot = doc.add_paragraph('Document version 1.0  ·  July 2026  ·  For support contact your Alphabot account manager.')
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.runs[0].font.size = Pt(8)
    foot.runs[0].font.color.rgb = C_GREY


# ══════════════════════════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════════════════════════

def main():
    out_path = os.path.join(os.path.dirname(__file__), 'Alphabot_User_Guide.docx')
    doc = setup_document()

    add_cover(doc)
    section_platform_manager(doc)
    section_platform_admin(doc)
    section_tenant_admin(doc)
    section_agent(doc)
    section_testing(doc)
    section_appendix(doc)

    doc.save(out_path)
    print(f'Saved: {out_path}')

if __name__ == '__main__':
    main()
