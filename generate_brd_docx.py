"""
BRD Generator — WhatsApp AI Agent Platform (.docx)
Generates a fully-editable Word document.
Run:  python generate_brd_docx.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "WhatsApp_AI_Agent_BRD.docx")

# ── Palette (RGB) ─────────────────────────────────────────────────────────────
DARK   = RGBColor(15,  23,  42)
BLUE   = RGBColor(37,  99, 235)
GRAY   = RGBColor(71,  85, 105)
LGRAY  = RGBColor(203, 213, 225)
WHITE  = RGBColor(255, 255, 255)
LBLUE  = RGBColor(219, 234, 254)


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_para_shading(para, hex_color: str):
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def set_para_border_left(para, hex_color="2563EB", size=24):
    pPr    = para._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    left   = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    str(size))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def add_run(para, text, bold=False, italic=False, color=None, size=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])


# ── Document styles ───────────────────────────────────────────────────────────

def style_document(doc: Document):
    style = doc.styles["Normal"]
    style.font.name  = "Calibri"
    style.font.size  = Pt(10)
    style.font.color.rgb = DARK

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)


# ── Cover page ────────────────────────────────────────────────────────────────

def cover(doc: Document):
    # Title block
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(title, "0F172A")
    r = title.add_run("WhatsApp AI Agent")
    r.bold = True
    r.font.size  = Pt(32)
    r.font.color.rgb = WHITE
    r.font.name  = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(sub, "0F172A")
    r2 = sub.add_run("Multi-Tenant SaaS Platform")
    r2.font.size  = Pt(14)
    r2.font.color.rgb = RGBColor(147, 197, 253)
    r2.font.name  = "Calibri"

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(banner, "2563EB")
    r3 = banner.add_run("Business Requirements Document")
    r3.bold = True
    r3.font.size  = Pt(18)
    r3.font.color.rgb = WHITE
    r3.font.name  = "Calibri"

    doc.add_paragraph()

    # Meta table
    meta = [
        ("Document Version", "1.0"),
        ("Status",           "Final Draft"),
        ("Date",             "July 2026"),
        ("Audience",         "Internal / Stakeholders"),
        ("Classification",   "Confidential"),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_col_widths(tbl, [5.5, 8.0])
    for i, (k, v) in enumerate(meta):
        kc, vc = tbl.rows[i].cells
        set_cell_bg(kc, "F1F5F9")
        set_cell_bg(vc, "F1F5F9")
        kp = kc.paragraphs[0]
        add_run(kp, k, bold=True, color=GRAY, size=10)
        vp = vc.paragraphs[0]
        add_run(vp, v, color=DARK, size=10)

    doc.add_paragraph()

    # Abstract
    abstract = doc.add_paragraph()
    abstract.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_abs = abstract.add_run(
        "This document defines the complete business requirements for the WhatsApp AI Agent "
        "platform — a multi-tenant SaaS product enabling businesses to deploy AI-powered "
        "WhatsApp bots and voice agents for customer support, sales engagement, and lifecycle "
        "management, with embedded payment collection and campaign tooling."
    )
    r_abs.italic = True
    r_abs.font.size  = Pt(10)
    r_abs.font.color.rgb = GRAY

    doc.add_page_break()


# ── Section helpers ───────────────────────────────────────────────────────────

def h1(doc: Document, text: str):
    p = doc.add_paragraph()
    set_para_border_left(p, "2563EB", 24)
    pf = p.paragraph_format
    pf.space_before = Pt(14)
    pf.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(16)
    r.font.color.rgb = DARK
    r.font.name  = "Calibri"
    return p


def h2(doc: Document, text: str):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(12)
    r.font.color.rgb = BLUE
    r.font.name  = "Calibri"
    # bottom border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2563EB")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def h3(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size  = Pt(10.5)
    r.font.color.rgb = DARK
    r.font.name  = "Calibri"
    return p


def body(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.size  = Pt(10)
    r.font.color.rgb = DARK
    r.font.name  = "Calibri"
    return p


def bullet(doc: Document, text: str, depth=0):
    style = "List Bullet" if depth == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5 + depth * 0.6)
    r = p.add_run(text)
    r.font.size  = Pt(10)
    r.font.color.rgb = DARK
    r.font.name  = "Calibri"
    return p


def callout(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.4)
    set_para_shading(p, "DBEAFE")
    set_para_border_left(p, "2563EB", 20)
    r = p.add_run(text)
    r.italic = True
    r.font.size  = Pt(9.5)
    r.font.color.rgb = GRAY
    r.font.name  = "Calibri"
    return p


def add_table(doc: Document, headers, rows, col_widths_cm):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = tbl.rows[0]
    for i, (h, cell) in enumerate(zip(headers, hdr_row.cells)):
        set_cell_bg(cell, "0F172A")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size  = Pt(9)
        r.font.color.rgb = WHITE
        r.font.name  = "Calibri"

    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        shade = "F8FAFB" if ri % 2 == 0 else "FFFFFF"
        for ci, (val, cell) in enumerate(zip(row_data, row.cells)):
            set_cell_bg(cell, shade)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size  = Pt(9)
            r.font.color.rgb = DARK
            r.font.name  = "Calibri"

    set_col_widths(tbl, col_widths_cm)
    doc.add_paragraph()
    return tbl


# ════════════════════════════════════════════════════════════════════════════
#  CONTENT
# ════════════════════════════════════════════════════════════════════════════

def content(doc: Document):

    # ── 1. Executive Summary ──────────────────────────────────────────────────
    h1(doc, "1.  Executive Summary")
    body(doc,
        "WhatsApp AI Agent is a B2B SaaS platform that enables businesses to deploy AI-powered "
        "conversational bots on WhatsApp and Voice channels without writing any code. The platform "
        "provides three specialised bot personas — Customer Support, Sales Engagement, and "
        "Lifecycle Management — each configurable through a no-code dashboard with layered "
        "guardrails, custom knowledge bases, escalation policies, voice call pipelines, campaign "
        "tools, and embedded payment collection."
    )
    body(doc,
        "Built for the Indian SMB and mid-market, the platform integrates natively with local "
        "payment gateways (PhonePe, Razorpay), local telephony (Exotel), and India's regional "
        "language AI providers (Sarvam AI). It supports 10+ regional languages and deploys "
        "on top of the Meta WhatsApp Business API."
    )
    h3(doc, "Core Value Propositions")
    for vp in [
        "Zero-code bot deployment — businesses configure, not code.",
        "Multi-bot, multi-tenant — one platform serves unlimited client organisations.",
        "4-layer guardrail cascade — platform → product → client → bot safety controls.",
        "Confidence-driven escalation — AI self-reports uncertainty and hands off to humans automatically.",
        "Omnichannel — WhatsApp messaging + outbound/inbound AI voice calls in one product.",
        "Embedded commerce — orders, UPI/card payment links, and delivery tracking inside conversations.",
        "India-first stack — Exotel, PhonePe, Razorpay, Sarvam AI, 10+ regional languages.",
        "Enterprise-grade security — Supabase Row-Level Security with per-tenant data isolation.",
    ]:
        bullet(doc, vp)

    # ── 2. Platform Architecture ──────────────────────────────────────────────
    h1(doc, "2.  Platform Architecture")
    h2(doc, "2.1  System Components")
    add_table(doc,
        ["Component", "Technology", "Responsibilities"],
        [
            ("API Server (apps/api)",       "Node.js / Fastify",    "Webhook processing, AI response pipeline, voice orchestration, campaign dispatch, KB ingest."),
            ("Web Application (apps/web)",  "Next.js 14 App Router","Dual portal: client dashboard + platform admin UI. Server Actions for all data mutations."),
            ("Shared Package",              "TypeScript",           "Cross-package types, enums, and shared utility functions."),
            ("Database Package",            "Supabase / PostgreSQL","Migrations, RLS policies, pgvector for embeddings, get_bot_context RPC."),
            ("Cache Layer",                 "Redis",                "KB lookup cache, bot context cache (60-second TTL), token quota counters."),
        ],
        [5.5, 3.8, 9.5]
    )

    h2(doc, "2.2  User Roles & Permissions")
    add_table(doc,
        ["Role", "Scope", "Permissions"],
        [
            ("Platform Manager", "SaaS Operator", "Full access: create/manage clients, products, guardrails, LLM configs, voice providers, billing, notifications, platform staff."),
            ("Platform Admin",   "SaaS Operator", "Read-only monitoring view of all client data and system health."),
            ("Client Manager",   "Tenant",        "Full workspace control: bot configs, guardrails, KB, team, WhatsApp numbers, LLM keys, orders."),
            ("Client Admin",     "Tenant",        "Escalation management, limited team management within the tenant."),
            ("Agent",            "Tenant",        "View and claim assigned escalations; send replies during human takeover."),
        ],
        [4.2, 3.5, 11.1]
    )

    h2(doc, "2.3  Tenant Plans")
    add_table(doc,
        ["Plan", "Conversation Limit / Month", "AI Token Quota"],
        [
            ("Starter", "500 conversations",   "Monthly cap enforced (plan-based)"),
            ("Growth",  "2,000 conversations", "Higher monthly cap (plan-based)"),
            ("Scale",   "Unlimited",           "Unlimited"),
        ],
        [3.5, 6.0, 9.3]
    )
    callout(doc,
        "Tenant status: trial → active (on subscription) or trial → expired. "
        "Suspended tenants receive no bot replies; existing conversation history is preserved."
    )

    # ── 3. Bot Types ──────────────────────────────────────────────────────────
    h1(doc, "3.  Bot Types (Products)")
    body(doc,
        "The platform ships three distinct bot products. Each can be activated independently "
        "per client with its own system prompt, knowledge base, guardrails, voice settings, "
        "LLM credentials, and escalation policy."
    )
    for slug, name, desc, caps in [
        ("support_bot", "Customer Support Bot",
         "Resolves customer queries by answering FAQs, looking up order/account status, "
         "and handling common support workflows via the knowledge base.",
         [
             "FAQ resolution from the knowledge base",
             "Order and account status lookup",
             "Confidence-based escalation to human agents",
             "CSAT rating collection after resolution (1–5 stars)",
             "Multilingual support (10+ Indian regional languages)",
         ]),
        ("sales_bot", "Sales Engagement Bot",
         "Qualifies inbound leads, shares product information, and detects buying intent "
         "to hand off warm prospects to a sales agent at the right moment.",
         [
             "Lead qualification through structured conversation stages",
             "Product information delivery from the KB",
             "Sales lead detection ([SALES_LEAD] marker) with auto-escalation to sales agent",
             "Customer entity capture (email, product interest, quantity, location)",
             "Conversation state machine: greeting → qualifying → resolving → following_up → closing",
         ]),
        ("lifecycle_bot", "Lifecycle Management Bot",
         "Handles post-sale engagement: order tracking, invoice queries, payment reminders, "
         "and renewal outreach integrated with the Orders and Payments module.",
         [
             "Order status and delivery tracking",
             "Invoice and payment query resolution",
             "Payment link dispatch (PhonePe / Razorpay) within the conversation",
             "Proactive follow-up reminders for pending payments",
             "Renewal and upsell outreach",
         ]),
    ]:
        h3(doc, f"{name}  —  slug: {slug}")
        body(doc, desc)
        for cap in caps:
            bullet(doc, cap, depth=1)

    # ── 4. Core Feature Modules ───────────────────────────────────────────────
    h1(doc, "4.  Core Feature Modules")

    h2(doc, "4.1  Inbound WhatsApp Message Pipeline")
    body(doc,
        "Every inbound WhatsApp message triggers the pipeline below. The HTTP 200 response "
        "is sent immediately; all AI processing is fully asynchronous."
    )
    for i, (title, desc) in enumerate([
        ("Provider Inference",        "Payload shape determines the WhatsApp provider (Meta Cloud vs Twilio). Falls back to the alternate provider if no active number is found for the inferred one."),
        ("Bot Context Load",          "Single get_bot_context RPC fetches: WhatsApp number credentials, tenant plan/status, bot config, all 4 guardrail layers, and LLM credentials. Result Redis-cached for 60 s."),
        ("Plan & Status Enforcement", "Suspended tenants silently dropped. Monthly conversation and token quotas checked before any AI work begins."),
        ("Contact Upsert",            "Contact identified by phone number (E.164) or BSUID (opaque WhatsApp username ID). Contact record and persistent memory JSON created or updated."),
        ("CSAT Intercept",            "If contact has an awaiting_csat flag and replies 1–5, the rating is recorded and a thank-you sent. No AI call is made."),
        ("Message Deduplication",     "whatsapp_msg_id unique constraint prevents duplicate AI calls when Meta retries webhooks."),
        ("Optimistic Processing Lock","30-second lock on the conversation row prevents concurrent AI calls from simultaneous retries."),
        ("Escalation Keyword Check",  "Configured trigger phrases checked before AI call. Match → conversation escalated, human-handover message sent immediately."),
        ("KB Retrieval (RAG)",        "4-strategy cascade: Redis cache → pgvector similarity search → keyword ILIKE search → product-scoped legacy fallback."),
        ("History Assembly",          "Up to 40 messages loaded. Recent messages passed verbatim; older turns compressed into an archive summary block in the system prompt."),
        ("AI Response Generation",    "Claude API called with assembled system prompt (guardrails, KB context, conversation stage, contact memory, language directive). Confidence score extracted."),
        ("Marker Extraction & Stripping","[SALES_LEAD], [STAGE:x], and [ENTITY:key=value] markers parsed and stripped before the message reaches the customer."),
        ("Guardrail Enforcement",     "Blocked keywords/topics checked across all 4 layers. Response truncated to the effective max_response_length (most restrictive layer wins)."),
        ("Escalation Policy",         "Low-confidence counter incremented. If consecutive low-confidence turns >= max_low_confidence_reprompts, conversation is auto-escalated."),
        ("Reply Dispatch",            "Message sent via provider gateway and stored with delivery_status=sent. Meta delivery/read receipts advance the status ladder: sent → delivered → read."),
        ("Lock Release",              "Processing lock cleared non-blocking. Conversation updated_at timestamp refreshed."),
    ], 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.3)
        add_run(p, f"{i}.  ", bold=True, color=BLUE, size=10)
        add_run(p, f"{title}:  ", bold=True, color=DARK, size=10)
        add_run(p, desc, color=DARK, size=10)

    h2(doc, "4.2  Conversation Management")
    h3(doc, "Conversation Statuses")
    add_table(doc,
        ["Status", "Description"],
        [
            ("open",       "Bot is actively responding to the customer."),
            ("escalated",  "Conversation handed to a human agent; bot is silent."),
            ("bot_paused", "Agent has taken manual control; messages stored but not auto-replied."),
            ("resolved",   "Closed by an agent or auto-resolved after inactivity."),
        ],
        [3.5, 15.3]
    )
    h3(doc, "AI Conversation State Machine")
    body(doc,
        "Conversations progress through structured stages. The AI declares stage transitions "
        "and captures customer entities by appending control markers at the end of its response "
        "(stripped before delivery):"
    )
    for stage in ["greeting", "qualifying", "resolving", "following_up", "closing"]:
        bullet(doc, stage)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    add_run(p, "Stage update:   ", bold=True, color=DARK, size=9.5)
    add_run(p, "[STAGE:qualifying]\n", color=GRAY, size=9.5)
    add_run(p, "Entity capture: ", bold=True, color=DARK, size=9.5)
    add_run(p, "[ENTITY:email=alice@example.com]  [ENTITY:order_id=ORD-123]", color=GRAY, size=9.5)

    h3(doc, "Delivery Status Ladder")
    body(doc,
        "Outbound messages move one-way: sent → delivered → read. "
        "The failed status is always accepted. Backward transitions are silently ignored."
    )

    h2(doc, "4.3  Knowledge Base (RAG)")
    body(doc,
        "Each tenant manages one or more KB Collections assigned to specific bots with a "
        "priority order. Documents are ingested, semantically chunked, and embedded via "
        "Voyage-3 into a pgvector index."
    )
    h3(doc, "Supported Document Formats")
    body(doc, "PDF, DOCX, TXT, Markdown, Images (OCR via image extraction)")
    h3(doc, "KB Entry Lifecycle")
    body(doc, "draft  →  review  →  live  →  archived")
    h3(doc, "4-Strategy RAG Fallback")
    for s in [
        "Redis cache — previously retrieved results cached per (tenant, product, query).",
        "Vector similarity search — pgvector cosine similarity using Voyage-3 embeddings.",
        "Keyword ILIKE search — full-text fallback when no vector match is found.",
        "Product-scoped legacy fallback — searches flat (non-collection) KB entries.",
    ]:
        bullet(doc, s)
    h3(doc, "KB-Only Mode")
    body(doc,
        "When enabled at any guardrail layer, the AI answers only from KB results. "
        "If no relevant entry is found it acknowledges the gap and offers to escalate. "
        "Configurable globally, per product type, per tenant, or per individual bot."
    )
    h3(doc, "AI-Generated KB Suggestions")
    body(doc,
        "Frequently asked questions not covered by the KB are flagged as suggestions "
        "(status: pending). Client managers review and promote them to live or dismiss them."
    )

    h2(doc, "4.4  Escalation Management")
    h3(doc, "Escalation Triggers")
    for t in [
        "Keyword match — customer types a configured phrase (e.g., 'speak to human', 'urgent', 'refund').",
        "Low-confidence consecutive turns — AI confidence below threshold for N turns in a row.",
        "Sales lead detected — [SALES_LEAD] marker triggers immediate handoff to a sales agent.",
        "KB-only no-match — KB-only mode active but no relevant KB entry was found.",
        "Auto-escalation timeout — conversation open for more than X configured hours.",
    ]:
        bullet(doc, t)
    h3(doc, "Escalation Flow")
    for s in [
        "Conversation status flipped to 'escalated'; bot goes silent.",
        "Escalation record created (status: pending).",
        "Email notification dispatched to client_manager and platform_admin via Brevo.",
        "If auto_dispatch_on_escalation is enabled: outbound voice call placed to the customer.",
        "Agent claims the escalation → conversation assigned (agent_id set on conversation).",
        "Agent sends manual WhatsApp replies until they mark the escalation resolved.",
        "On resolution: conversation closed; optional CSAT prompt sent to the customer.",
    ]:
        bullet(doc, s)
    h3(doc, "Human Takeover (bot_paused)")
    body(doc,
        "An agent can pause the bot without a full escalation — useful for brief interventions. "
        "Inbound messages are stored but not auto-replied. The agent can resume the bot at "
        "any time by switching the status back to 'open'."
    )
    h3(doc, "CSAT Collection")
    body(doc,
        "After resolution the bot sends a 1–5 star rating prompt. The score is stored in "
        "contact.memory_json with a timestamp. A confirmation message is sent to the customer."
    )

    h2(doc, "4.5  Voice Call Module")
    body(doc,
        "The voice module enables AI-powered outbound phone calls via a configurable "
        "three-component stack: Telephony + Speech-to-Text + Text-to-Speech. "
        "Each call runs a multi-turn AI conversation loop with full transcript storage "
        "and automated structured outcome extraction."
    )
    h3(doc, "Supported Voice Providers")
    add_table(doc,
        ["Component", "Providers", "Notes"],
        [
            ("Telephony", "Twilio, Exotel",                              "Exotel recommended in India (lower per-minute cost). Twilio for international."),
            ("STT",       "Deepgram Nova-2, Sarvam AI, Azure Speech",    "Sarvam recommended for Hindi and Indian regional languages."),
            ("TTS",       "Twilio Say (free), Google Neural2, Sarvam TTS, Exotel Say", "Twilio Say is zero-cost built-in; Google/Sarvam for higher quality."),
        ],
        [3.2, 6.5, 9.1]
    )
    h3(doc, "Voice Call Flow")
    for step in [
        "Dispatch request (API, campaign trigger, or escalation auto-dispatch) creates a voice_calls record.",
        "Telephony provider places the outbound call.",
        "On answer: TwiML/ExoML greeting message played; customer prompted to speak.",
        "Customer audio → STT transcription → Claude AI response → TTS audio returned as next prompt.",
        "Multi-turn loop repeats until max_turns reached, silence timeout, or customer hangs up.",
        "Final status received from provider: completed / failed / voicemail / no_answer / busy.",
        "Claude analyses the full transcript and extracts structured outcome JSON.",
        "Call cost in INR calculated and stored alongside the full transcript.",
    ]:
        bullet(doc, step)
    h3(doc, "Post-Call Outcome Extraction")
    body(doc, "Claude analyses the transcript and returns:")
    for f in ["intent", "product_interest", "resolved (boolean)", "escalation_needed (boolean)", "sentiment", "follow_up_action", "summary"]:
        bullet(doc, f, depth=1)
    h3(doc, "Voice Configuration Options (per-bot)")
    for cfg in [
        "telephony_provider, stt_provider, tts_provider — override platform-level defaults",
        "language — en-IN, hi-IN, ta-IN, te-IN, kn-IN, mr-IN, gu-IN, ml-IN, en-US",
        "tts_voice — provider-specific voice name (e.g., Polly.Aditi, en-IN-Neural2-A)",
        "max_call_duration_seconds — maximum call length (default: 300 s)",
        "max_turns — maximum AI conversation turns per call (default: 20)",
        "silence_timeout_seconds — silence before hang-up (default: 5 s)",
        "voicemail_enabled / voicemail_message — leave a voicemail on no-answer",
        "greeting_message — opening message (overrides the product default)",
        "auto_dispatch_on_escalation — auto-call customer when their conversation is escalated",
        "escalation_voice_delay_seconds — delay before dispatching the escalation call",
    ]:
        bullet(doc, cfg)

    h2(doc, "4.6  Campaign Engine")
    body(doc,
        "Campaigns enable bulk outbound engagement via WhatsApp, Voice, or both simultaneously. "
        "Per-contact delivery and response status is tracked with configurable retry logic."
    )
    h3(doc, "Campaign Channels")
    add_table(doc,
        ["Channel", "Behaviour"],
        [
            ("WhatsApp", "Sends a WhatsApp template message to all contacts. Tracks: sent, replied, failed per contact."),
            ("Voice",    "Dispatches outbound AI calls to all contacts. Tracks: initiated, answered, voicemail, failed."),
            ("Both",     "Sends WhatsApp first; if no reply after configured delay, dispatches a voice call."),
        ],
        [3.0, 15.8]
    )
    h3(doc, "Campaign Lifecycle")
    body(doc, "draft  →  active (on launch)  →  paused (manual)  →  completed / cancelled")
    h3(doc, "Retry Logic")
    body(doc,
        "Per-campaign retry_config: retry_after_hours (wait before retry) and max_retries "
        "(maximum attempts per contact). Contacts in failed state are retried within the window."
    )
    h3(doc, "Real-Time Campaign Stats")
    body(doc,
        "stats_json tracks: total contacts, WhatsApp sent/replied/failed, "
        "voice answered/voicemail/failed — all displayed on the campaign detail page."
    )

    h2(doc, "4.7  Follow-Up Automation")
    body(doc, "Proactively re-engages contacts whose conversation has been idle for a configurable number of days.")
    h3(doc, "Configuration Options")
    for cfg in [
        "enabled — toggle on/off per bot type",
        "idle_days — days of inactivity before follow-up triggers",
        "message_template — the follow-up message text",
        "max_follow_ups — maximum follow-ups per conversation before stopping",
        "contact_scope — all contacts or a specific segment",
    ]:
        bullet(doc, cfg)
    h3(doc, "Follow-Up Statuses")
    body(doc, "scheduled  →  sent / failed / cancelled")

    h2(doc, "4.8  Orders & Payment Collection")
    body(doc,
        "Agents or the lifecycle bot can create orders within a conversation. Payment links "
        "(PhonePe / Razorpay) are generated and sent via WhatsApp. Confirmations arrive via webhook."
    )
    h3(doc, "Order Workflow")
    for step in [
        "Create order: add line items (name, quantity, unit price) and set total.",
        "Select payment provider and generate a payment link.",
        "Payment link sent to customer via WhatsApp.",
        "Customer pays → provider webhook received at /api/payments/{provider}/webhook.",
        "Payment status updated to 'paid'; order status updated to 'confirmed'.",
        "Conversation continues — bot can confirm delivery and handle post-payment queries.",
    ]:
        bullet(doc, step)
    h3(doc, "Payment Providers")
    for p_name, d in [
        ("PhonePe",  "UPI-based payments for India. SHA256-signed payloads. Redirect to PhonePe checkout."),
        ("Razorpay", "Multi-method: UPI, cards, net banking. Order API + webhook-driven confirmation."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        add_run(p, p_name + ":  ", bold=True, color=DARK, size=10)
        add_run(p, d, color=DARK, size=10)
    h3(doc, "Order Statuses")
    body(doc, "pending  →  confirmed  →  dispatched  →  delivered  →  cancelled")

    # ── 5. Configuration System ───────────────────────────────────────────────
    h1(doc, "5.  Configuration System")

    h2(doc, "5.1  4-Layer Guardrail Cascade")
    body(doc,
        "Guardrails control what the AI can say and how it handles sensitive topics. "
        "Four layers are merged at runtime using these merge rules:"
    )
    for rule in [
        "Lists (blocked_topics, blocked_keywords): UNION — all layers contribute their lists.",
        "Numeric limits (max_response_length): MIN — most restrictive layer wins.",
        "Booleans (kb_only_mode, no_external_links, no_personal_data): OR — any layer can enable.",
        "Action/message strings (on_blocked_topic, custom_blocked_message): most specific layer wins.",
    ]:
        bullet(doc, rule)
    add_table(doc,
        ["#", "Layer", "Set By", "Controls"],
        [
            ("1", "Global Platform Settings", "Platform Manager",
             "global_blocked_topics, global_blocked_keywords, max_response_length (default 2000), enforce_kb_only_globally, no_personal_data, no_external_links"),
            ("2", "Bot-Type Guardrails", "Platform Manager",
             "Per product slug: blocked_topics, blocked_keywords, max_response_length, kb_only_mode, no_personal_data, no_external_links, on_blocked_topic"),
            ("3", "Tenant Guardrails", "Client Manager",
             "Shared across all bots for the tenant: same as Layer 2 plus custom_blocked_message"),
            ("4", "Bot Config Guardrails", "Client Manager",
             "Per-bot overrides: tone, no_phone_numbers_in_response, on_low_confidence, reprompt_message, custom_blocked_message, max_response_length (default 1000)"),
        ],
        [0.8, 4.5, 4.0, 9.5]
    )

    h2(doc, "5.2  Bot Configuration Settings")
    h3(doc, "Core Settings")
    for cfg in [
        "system_prompt — AI persona and business-specific instructions (overrides product default)",
        "ai_model — Anthropic model ID (e.g., claude-sonnet-4-6); overrides product default",
        "confidence_threshold — float 0–1; below this triggers low-confidence reprompt (default 0.6)",
        "escalation_triggers — keyword phrases for immediate human escalation",
        "kb_only_mode — restrict AI to KB answers only",
    ]:
        bullet(doc, cfg)
    h3(doc, "Escalation Policy")
    for cfg in [
        "confidence_threshold — per-bot override of the platform default",
        "max_low_confidence_reprompts — consecutive low-confidence turns before escalating (default 2)",
        "on_exhaust — 'escalate' or 'silent' when reprompt limit is reached",
        "reprompt_message — optional message sent to the customer before the escalation turn",
        "auto_escalate_after_hours — auto-escalate after N hours open without resolution (null = off)",
    ]:
        bullet(doc, cfg)
    h3(doc, "Tone & Content Filters")
    for cfg in [
        "tone — professional | casual | empathetic | formal (injected into system prompt)",
        "no_external_links — prevent the AI from including URLs",
        "no_phone_numbers_in_response — prevent the AI from sharing phone numbers",
        "no_personal_data — prevent the AI from referencing PII",
        "on_blocked_topic — escalate | silent | custom_message",
        "custom_blocked_message — text shown when a blocked topic is detected",
    ]:
        bullet(doc, cfg)

    h2(doc, "5.3  LLM Configuration Hierarchy (6 Levels)")
    body(doc,
        "The platform resolves which AI model and API credentials to use through a 6-level "
        "hierarchy. The first valid match from the top wins:"
    )
    add_table(doc,
        ["Level", "Source"],
        [
            ("1 — Client + Bot-specific",   "llm_configs row with matching tenant_id AND product_slug"),
            ("2 — Client generic",           "llm_configs row with matching tenant_id, no product_slug"),
            ("3 — Platform + Bot-specific",  "llm_configs row with no tenant_id, matching product_slug"),
            ("4 — Platform generic",         "llm_configs row with no tenant_id, no product_slug"),
            ("5 — Bot config model",         "bot_configs.ai_model (model only; uses platform API key)"),
            ("6 — Product default",          "products.default_model (product catalog; uses platform API key)"),
        ],
        [5.5, 13.3]
    )
    callout(doc,
        "Auth-failure fallback: if a client's custom API key returns a 401/403 error, the platform "
        "automatically retries the AI call using its own platform key before returning an error to the customer."
    )

    h2(doc, "5.4  WhatsApp Number Management")
    add_table(doc,
        ["Provider", "Configuration Required"],
        [
            ("Meta Cloud API", "phone_number_id, access_token, verify_token, app_secret — official WhatsApp Business API."),
            ("Twilio",         "Account SID, Auth Token, WhatsApp-enabled phone number."),
            ("Interakt",       "API key and endpoint — third-party BSP."),
            ("WATI",           "API token and endpoint — third-party BSP."),
            ("Gupshup",        "API key and source phone — third-party BSP."),
        ],
        [4.0, 14.8]
    )

    # ── 6. Analytics & Reporting ──────────────────────────────────────────────
    h1(doc, "6.  Analytics & Reporting")
    h2(doc, "6.1  Client Dashboard Analytics  (/analytics)")
    body(doc, "Available with 7-day and 30-day toggles.")
    for m in [
        "Total, open, escalated, resolved, and bot_paused conversation counts",
        "Daily message volume trend (bar chart)",
        "Monthly AI token consumption rolling total",
        "Escalation rate and average resolution time",
        "Per-product breakdown (support_bot, sales_bot, lifecycle_bot)",
        "Contact sentiment distribution (positive, neutral, negative, frustrated)",
    ]:
        bullet(doc, m)
    h2(doc, "6.2  Platform-Level Analytics  (/platform/analytics)")
    for m in [
        "Total active clients and active trials",
        "Platform-wide conversation volume across all tenants",
        "Trial utilisation rates and paid conversion trends",
        "Per-client breakdown of message and token usage",
    ]:
        bullet(doc, m)
    h2(doc, "6.3  Usage Events (Metered Billing Data)")
    add_table(doc,
        ["Event Type", "Triggered When"],
        [
            ("conversation_started", "A new conversation is created for a contact."),
            ("message_sent",         "Each inbound message or bot reply."),
            ("ai_token_used",        "An AI response is generated (input + output token count stored)."),
            ("escalation",           "A conversation is escalated to a human agent."),
            ("kb_query",             "A KB lookup is performed for an incoming message."),
        ],
        [5.5, 13.3]
    )

    # ── 7. Notification System ────────────────────────────────────────────────
    h1(doc, "7.  Notification System")
    body(doc,
        "Event-driven email notifications sent via Brevo (Sendinblue). Rules are configurable "
        "per event at platform-wide or per-tenant scope with role-based or custom email targeting."
    )
    add_table(doc,
        ["Event", "Description"],
        [
            ("trial_expiring_7d",    "Trial expires in 7 days — alert client manager."),
            ("trial_expiring_1d",    "Trial expires tomorrow — alert client manager + platform manager."),
            ("trial_expired",        "Trial ended — notify both parties."),
            ("client_invited",       "New client invite sent via the platform."),
            ("client_activated",     "Client completed onboarding and went live."),
            ("escalation_created",   "New escalation — notify assigned agent and managers."),
            ("escalation_timeout",   "Escalation unclaimed past the threshold duration."),
            ("new_client_onboarded", "Platform manager alert on new client signup."),
            ("daily_report",         "Daily digest: conversation count, escalation rate, token usage."),
            ("low_confidence_spike", "Unusual spike in low-confidence AI responses detected."),
            ("bot_error",            "Runtime error in the bot pipeline."),
            ("subscription_renewed", "Billing cycle renewed successfully."),
        ],
        [5.5, 13.3]
    )
    h3(doc, "Recipient Targeting")
    for t in [
        "Role-based: platform_admin, platform_manager, client_manager, client_admin",
        "Custom email addresses (comma-separated list)",
        "Scope: platform-wide (all tenants) or per-tenant (specific client only)",
    ]:
        bullet(doc, t)

    # ── 8. Free Trial Management ──────────────────────────────────────────────
    h1(doc, "8.  Free Trial Management")
    body(doc,
        "Platform managers can grant time-limited free trials per product per client. "
        "Trials operate independently from paid subscriptions."
    )
    for item in [
        "Duration: configurable per trial (default 14 days).",
        "Scope: per product — clients can trial each bot type independently.",
        "Model restriction: trial may be limited to a specific AI model.",
        "Status flow:  active  →  expired  or  active  →  converted (on subscription).",
        "Automatic notifications: 7-day warning, 1-day warning, and expiry notification.",
        "Conversion tracking: flagged when a trial converts to a paid subscription.",
        "Trial utilisation tracked in platform analytics for conversion rate reporting.",
    ]:
        bullet(doc, item)

    # ── 9. Team Management ────────────────────────────────────────────────────
    h1(doc, "9.  Team Management")
    h2(doc, "9.1  Client Team")
    body(doc, "Client managers invite team members via email. Invite links expire in 7 days.")
    add_table(doc,
        ["Role", "Permissions"],
        [
            ("client_manager", "Full workspace control: bot configs, guardrails, KB, team, WhatsApp numbers, LLM keys, orders, billing."),
            ("client_admin",   "Escalation management, limited team management within the tenant."),
            ("agent",          "View/claim assigned escalations; send WhatsApp replies during human takeover."),
        ],
        [4.0, 14.8]
    )
    h2(doc, "9.2  Platform Team")
    body(doc, "Platform managers invite internal SaaS operator staff.")
    add_table(doc,
        ["Role", "Permissions"],
        [
            ("manager", "Full platform control: create and manage all clients, products, configs."),
            ("admin",   "Read-only monitoring access to all client data for support and investigation."),
        ],
        [4.0, 14.8]
    )

    # ── 10. Third-Party Integrations ──────────────────────────────────────────
    h1(doc, "10.  Third-Party Integrations")
    add_table(doc,
        ["Service", "Category", "Usage"],
        [
            ("Meta Cloud API",            "WhatsApp",          "Primary WhatsApp Business API. Webhook verification, message send, delivery receipts."),
            ("Twilio",                    "WhatsApp + Voice",  "WhatsApp Sandbox + outbound/inbound AI voice calls (TwiML)."),
            ("Interakt / WATI / Gupshup", "WhatsApp BSP",     "Alternative third-party WhatsApp providers."),
            ("Exotel",                    "Telephony (India)", "India-focused outbound voice (ExoML). Lower per-minute cost than Twilio."),
            ("Deepgram Nova-2",           "STT",               "Fast, accurate English speech-to-text transcription."),
            ("Sarvam AI",                 "STT + TTS",         "Indian regional language support: hi-IN, ta-IN, te-IN, kn-IN, mr-IN, gu-IN, ml-IN."),
            ("Azure Speech",              "STT",               "Enterprise-grade STT (optional)."),
            ("Google Neural2 TTS",        "TTS",               "High-quality TTS for English and Indian regional languages."),
            ("Twilio Say",                "TTS",               "Free built-in TTS using Amazon Polly voices (Aditi, Raveena)."),
            ("PhonePe",                   "Payments",          "UPI payment links. SHA256-signed payloads. India-first."),
            ("Razorpay",                  "Payments",          "Multi-method payment links (UPI, cards, net banking). Webhook-driven."),
            ("Anthropic / Claude",        "AI / LLM",          "Primary AI for bot responses and post-call outcome extraction."),
            ("OpenRouter",                "AI / LLM",          "Multi-model LLM routing option."),
            ("OpenAI",                    "AI / LLM",          "Direct OpenAI models as alternative LLM provider."),
            ("Voyage-3",                  "Embeddings",        "Semantic KB search embeddings stored in pgvector."),
            ("Supabase",                  "Database",          "PostgreSQL + pgvector + Row-Level Security + Auth."),
            ("Redis",                     "Cache",             "KB lookup cache, bot context cache (60 s TTL), token counters."),
            ("Brevo (Sendinblue)",        "Email",             "Transactional and event-driven notification emails."),
        ],
        [4.5, 3.8, 10.5]
    )

    # ── 11. Core Data Model ───────────────────────────────────────────────────
    h1(doc, "11.  Core Data Model")
    for group, items in [
        ("Identity & Access", [
            "tenants — Workspace/organisation (plan, status, provider).",
            "tenant_users — Team members linked to a tenant (role, user_id).",
            "platform_users — Platform staff (role: manager | admin).",
            "client_invites — Invite tokens (email, role, expires in 7 days).",
        ]),
        ("Communication", [
            "whatsapp_numbers — Business numbers (provider, product_slug, config_json, active flag).",
            "contacts — Customer records (phone/BSUID, name, memory_json: preferences, sentiment, CSAT).",
            "conversations — Instances (status, product_type, stage, ai_vars, assigned_agent_id, processing_lock).",
            "messages — History (role, content, media_url, confidence_score, delivery_status, whatsapp_msg_id).",
            "escalations — Records (trigger_reason, agent_id, status: pending | assigned | resolved).",
            "agent_sessions — Human takeover periods (started_at, ended_at, resolution_note).",
        ]),
        ("AI & Configuration", [
            "products — Product catalog (slug, name, default_prompt, default_model, active).",
            "bot_configs — Per-bot settings (system_prompt, confidence_threshold, guardrails_json, voice_config).",
            "platform_settings — Global key-value store for guardrails and feature flags.",
            "bot_type_guardrails — Layer 2 guardrails scoped to a product_slug.",
            "tenant_guardrails — Layer 3 guardrails shared across all bots for a tenant.",
            "llm_configs — LLM credentials (tenant_id x product_slug hierarchy; api_key, model).",
        ]),
        ("Knowledge Base", [
            "kb_collections — Named KB collections (tenant, name, description).",
            "kb_collection_bots — Assigns collections to bots with priority ordering.",
            "knowledge_base — KB entries (question, answer, category, embedding vector, status).",
            "kb_documents — Uploaded documents (PDF/DOCX/TXT; processing status; chunk_count).",
            "kb_suggestions — AI-generated FAQ suggestions (frequency count, status).",
        ]),
        ("Commerce", [
            "orders — Customer orders (items_json, total, status, contact_id, conversation_id).",
            "payments — Payment records (order_id, provider, status, link_url, webhook_received_at).",
            "campaigns — Outbound campaigns (channel, template_id, status, stats_json, retry_config).",
            "campaign_contacts — Per-contact tracking (whatsapp_status, voice_status, attempts).",
            "follow_up_configs — Follow-up settings (idle_days, max_follow_ups, message_template).",
            "follow_up_sequences — Scheduled follow-up send log.",
        ]),
        ("Voice & Telephony", [
            "voice_provider_configs — Provider registry (component: telephony | stt | tts, credentials, cost_per_min_inr).",
            "voice_calls — Call log (status, duration_seconds, transcript, outcome_json, cost_rupees, triggered_by).",
        ]),
        ("Billing & Usage", [
            "free_trials — Time-limited trials (tenant_id, product_slug, starts_at, ends_at, status).",
            "subscriptions — Billing records (product, tier, billing_cycle, next_billing_date).",
            "usage_events — Metered usage rows (event_type, token_count, timestamp).",
            "notification_configs — Event-triggered email rules (scope, event_type, recipients, enabled).",
            "tenant_products — Product activation per tenant (tier: base | advanced).",
        ]),
    ]:
        h3(doc, group)
        for item in items:
            bullet(doc, item)

    # ── 12. Security & Data Isolation ─────────────────────────────────────────
    h1(doc, "12.  Security & Data Isolation")
    h2(doc, "12.1  Row-Level Security (RLS)")
    body(doc,
        "Supabase RLS policies enforce strict per-tenant data isolation at the database level. "
        "All web app queries run under authenticated user context; the API server uses the "
        "service role only for webhook processing."
    )
    for item in [
        "Tenants can only read/write their own conversations, contacts, messages, and KB.",
        "Agents see only conversations/escalations belonging to their tenant.",
        "Platform staff: admin (read-all) vs manager (read-write-all) enforced by policy.",
        "get_user_tenant_id() Postgres function resolves the caller's tenant from JWT claims.",
    ]:
        bullet(doc, item)
    h2(doc, "12.2  Credential & Secret Management")
    for item in [
        "WhatsApp credentials (access tokens, phone number IDs) stored encrypted in config_json per number.",
        "Payment credentials (merchant IDs, salt keys) held in environment variables only.",
        "LLM API keys stored in llm_configs; masked in all API responses.",
        "Voice provider credentials in voice_provider_configs.credentials_json.",
        "Webhook signatures validated (Meta HMAC app_secret, Razorpay webhook secret).",
        "Environment variables managed via deployment platform — never committed to source control.",
    ]:
        bullet(doc, item)
    h2(doc, "12.3  Anti-Abuse Controls")
    for item in [
        "30-second optimistic lock prevents duplicate AI calls from Meta webhook retries.",
        "whatsapp_msg_id unique constraint prevents duplicate message storage.",
        "Plan limits (conversation count + monthly token quota) checked before every AI call.",
        "Suspended tenant check at webhook entry — no processing for suspended accounts.",
        "Blocked keyword/topic guardrails enforced at every layer before sending any AI response.",
    ]:
        bullet(doc, item)

    # ── 13. Non-Functional Requirements ──────────────────────────────────────
    h1(doc, "13.  Non-Functional Requirements")
    h2(doc, "13.1  Performance")
    for item in [
        "Webhook HTTP response < 100 ms (200 sent immediately; AI work is fully async).",
        "Bot context Redis cache (60 s TTL) reduces Supabase round-trips on hot paths.",
        "KB lookup target < 300 ms (Redis cache hit < 5 ms; vector search < 200 ms).",
        "get_bot_context RPC consolidates 7 Supabase queries into a single database call.",
    ]:
        bullet(doc, item)
    h2(doc, "13.2  Reliability & Resilience")
    for item in [
        "Meta webhook retries handled via optimistic lock + whatsapp_msg_id deduplication.",
        "AI API key fallback: client key auth failure → automatic retry with platform key.",
        "Provider inference fallback: wrong WhatsApp provider inferred → retry with alternate.",
        "Graceful degradation: if AI pipeline fails, customer receives a human-escalation message.",
        "Voice call status webhooks handle all terminal states (completed, failed, voicemail, no_answer, busy).",
    ]:
        bullet(doc, item)
    h2(doc, "13.3  Scalability")
    for item in [
        "Stateless API server — horizontal scaling supported.",
        "Redis caching for high-frequency bot context and KB lookup paths.",
        "pgvector for semantic KB search at large embedding volumes.",
        "Metered usage events provide real-time billing and capacity planning data.",
    ]:
        bullet(doc, item)
    h2(doc, "13.4  Internationalisation")
    for item in [
        "Auto language detection: Hindi (Devanagari), Tamil, Telugu, Kannada, Punjabi, Gujarati, Odia, Arabic, Chinese.",
        "Bot auto-replies in the detected language/script when non-Latin text is received.",
        "Voice STT/TTS supports 8+ Indian regional language codes (en-IN, hi-IN, ta-IN, te-IN, etc.).",
        "Payment providers (PhonePe, Razorpay) and telephony (Exotel) optimised for India.",
    ]:
        bullet(doc, item)

    # ── 14. Glossary ──────────────────────────────────────────────────────────
    h1(doc, "14.  Glossary")
    add_table(doc,
        ["Term", "Definition"],
        [
            ("BSUID",           "WhatsApp Business Unique Identifier — opaque ID used instead of a phone number for users with a WhatsApp username."),
            ("Bot Context",     "Resolved configuration (system prompt, guardrails, LLM credentials, voice config) loaded per-request from Redis or RPC."),
            ("BSP",             "Business Solution Provider — third-party WhatsApp API provider (Interakt, WATI, Gupshup)."),
            ("CSAT",            "Customer Satisfaction Score — 1–5 star rating collected after conversation resolution."),
            ("ExoML",           "Exotel Markup Language — XML format for controlling Exotel voice call flows."),
            ("Guardrail",       "A safety rule applied to AI responses: blocked keywords, topic restrictions, content filters, response length limits."),
            ("KB",              "Knowledge Base — structured repository of Q&A entries and documents used for RAG-based AI augmentation."),
            ("LLM",             "Large Language Model — the AI model that generates bot responses (default: Claude by Anthropic)."),
            ("Optimistic Lock", "30-second database lock on a conversation row to prevent concurrent AI calls from simultaneous webhook retries."),
            ("Product Slug",    "String identifier for a bot type: support_bot, sales_bot, or lifecycle_bot."),
            ("RAG",             "Retrieval-Augmented Generation — injecting relevant KB results into the AI prompt before generating a response."),
            ("RLS",             "Row-Level Security — Supabase/PostgreSQL feature enforcing per-tenant data isolation at the query level."),
            ("STT",             "Speech-to-Text — converts customer voice audio to text for AI processing during voice calls."),
            ("Tenant",          "A client organisation using the platform; has its own team, bots, contacts, conversations, and billing."),
            ("TTS",             "Text-to-Speech — converts the AI's text response to audio delivered during voice calls."),
            ("TwiML",           "Twilio Markup Language — XML format for controlling Twilio voice call flows."),
        ],
        [3.5, 15.3]
    )

    # Closing
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(p, "DBEAFE")
    r = p.add_run("End of Document  —  WhatsApp AI Agent BRD v1.0  —  July 2026")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY


# ════════════════════════════════════════════════════════════════════════════
#  MAIN
# ════════════════════════════════════════════════════════════════════════════

def main():
    doc = Document()
    style_document(doc)

    cover(doc)
    content(doc)

    doc.save(OUTPUT)
    print(f"\nBRD (Word) saved to:\n  {OUTPUT}\n")


if __name__ == "__main__":
    main()
